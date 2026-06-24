from pathlib import Path
import os, shutil, zipfile, textwrap, math
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/mnt/data/qe2022_applied_lung_cancer')
if ROOT.exists(): shutil.rmtree(ROOT)
for p in [
    ROOT/'data/raw', ROOT/'data/dictionary', ROOT/'scripts', ROOT/'docs/source',
    ROOT/'docs/general_reference', ROOT/'output/tables', ROOT/'output/figures',
    ROOT/'output/appendix', ROOT/'output/word'
]: p.mkdir(parents=True, exist_ok=True)

# ---------------- Synthetic data ----------------
rng = np.random.default_rng(20220729)
n = 130
ID = np.arange(1, n+1)
age = np.clip(rng.normal(66, 9, n), 42, 86).round(0)
sex = rng.choice(['Female','Male'], n, p=[0.46,0.54])
smoking = rng.choice(['Never','Former','Current'], n, p=[0.18,0.62,0.20])
resp = rng.binomial(1, 1/(1+np.exp(-(-2.1 + .025*(age-60) + .75*(smoking!='Never')))))
card = rng.binomial(1, 1/(1+np.exp(-(-2.0 + .035*(age-60)))))
stroke = rng.binomial(1, 0.055 + 0.002*np.maximum(age-65,0))
ckd = rng.binomial(1, 1/(1+np.exp(-(-2.7 + .035*(age-60)))))
dm = rng.binomial(1, 1/(1+np.exp(-(-1.9 + .025*(age-60)))))
oth = rng.binomial(1, 0.22, n)
stage = rng.choice(['I','II','III'], n, p=[0.50,0.29,0.21])
histology = rng.choice(['Adenocarcinoma','Squamous cell carcinoma','Other'], n, p=[0.58,0.31,0.11])
tumor = np.clip(rng.gamma(2.5, 1.25, n) + np.where(stage=='III',1.0,0), .5, 9.5)
hilar = rng.binomial(1, np.where(stage=='I',.12,np.where(stage=='II',.38,.62)))
medi = rng.binomial(1, np.where(stage=='I',.04,np.where(stage=='II',.18,.50)))
surgery = []
for s, t in zip(stage, tumor):
    probs = [0.66,0.27,0.07] if s=='I' and t<4 else ([0.55,0.25,0.20] if s=='II' else [0.46,0.18,0.36])
    surgery.append(rng.choice(['Lobectomy','Wedge/segmentectomy','Pneumonectomy'], p=probs))
surgery=np.array(surgery)

# correlated PFTs
fev_pct = 88 - .45*(age-60) - 12*resp - 5*(smoking=='Current') - 3*(smoking=='Former') + rng.normal(0,10,n)
fev_pct = np.clip(fev_pct, 35, 125)
fev1 = 0.028*fev_pct + np.where(sex=='Male',.35,0) + rng.normal(0,.28,n)
fev1 = np.clip(fev1, .65, 4.2)
dlco_pct = 85 - .32*(age-60) - 10*resp - 8*(smoking=='Current') - 4*(smoking=='Former') + rng.normal(0,11,n)
dlco_pct = np.clip(dlco_pct, 30, 125)
dlco = .22*dlco_pct + np.where(sex=='Male',2.5,0) + rng.normal(0,2.2,n)
dlco = np.clip(dlco, 5, 35)

# complications
surg_effect = np.select([surgery=='Pneumonectomy', surgery=='Lobectomy'], [1.0,.35], default=0)
lin_any = -2.2 + .022*(age-65) + .85*resp + .55*card + .75*surg_effect - .015*(dlco_pct-75)
p_any = 1/(1+np.exp(-lin_any))
anycomp = rng.binomial(1,p_any)
respcomp = rng.binomial(1, np.clip(.02 + .28*anycomp + .12*resp + .08*(surgery=='Pneumonectomy'),0,.85))
cardcomp = rng.binomial(1, np.clip(.01 + .15*anycomp + .10*card,0,.65))
infectcomp = rng.binomial(1, np.clip(.025 + .20*anycomp + .06*dm,0,.55))
deathcomp = rng.binomial(1, np.clip(.005 + .04*anycomp + .025*(surgery=='Pneumonectomy'),0,.18))
# Ensure anycomp coherent with components
anycomp = np.maximum(anycomp, np.maximum.reduce([respcomp,cardcomp,infectcomp,deathcomp]))

# LOS
log_los = np.log(4.8) + .45*anycomp + .26*(surgery=='Pneumonectomy') + .09*(surgery=='Lobectomy') + .006*(age-65) - .004*(dlco_pct-75) + rng.normal(0,.35,n)
los = np.clip(np.exp(log_los),1,35).round(0)

# Survival
stage_lp = np.select([stage=='II',stage=='III'], [.48,1.05], default=0)
lp = .025*(age-65) + stage_lp + .10*(tumor-3) + .50*medi + .22*hilar - .012*(dlco_pct-75) + .20*card
base_haz=.0105
event_time = rng.exponential(1/(base_haz*np.exp(lp)))
censor_time = rng.uniform(18,105,n)
osmonths = np.minimum(event_time,censor_time)
died = (event_time<=censor_time).astype(int)
osmonths = np.clip(osmonths, .2, 120).round(1)

# Preop PFT indicator and missingness
preop = np.array(['Yes']*n, dtype=object)
df = pd.DataFrame({
    'ID':ID,'Age':age.astype(float),'Sex':sex,'Smoking':smoking,'Surgical_resection':'Yes',
    'type_of_surgery':surgery,'pathologic_stage':stage,'histology':histology,
    'Tumor_size_cm':tumor.round(1),'Hilar_LN_involved':np.where(hilar==1,'Yes','No'),
    'Mediastinal_LN_involved':np.where(medi==1,'Yes','No'),'Pre_op_PFTS':preop,
    'FEV1':fev1.round(2),'FEV_percent_OF_PREDICTED':fev_pct.round(1),
    'DLCO':dlco.round(1),'DLCO_percent_of_Predicted':dlco_pct.round(1),
    'Post_op_complications':np.where(anycomp==1,'Yes','No'),
    'Length_of_Hospital_stay':los.astype(float),
    'respiratory':np.where(resp==1,'Yes','No'),'cardiac':np.where(card==1,'Yes','No'),
    'stroke':np.where(stroke==1,'Yes','No'),'CKD':np.where(ckd==1,'Yes','No'),
    'DM':np.where(dm==1,'Yes','No'),'othcomorb':np.where(oth==1,'Yes','No'),
    'respcomp':np.where(respcomp==1,'Yes','No'),'cardcomp':np.where(cardcomp==1,'Yes','No'),
    'deathcomp':np.where(deathcomp==1,'Yes','No'),'infectcomp':np.where(infectcomp==1,'Yes','No'),
    'anycomp':np.where(anycomp==1,'Yes','No'),'died':died,'oasmons':osmonths
})
# Add controlled missingness
for col, frac in {
    'FEV1':.06,'FEV_percent_OF_PREDICTED':.055,'DLCO':.075,'DLCO_percent_of_Predicted':.07,
    'Tumor_size_cm':.025,'Smoking':.015,'histology':.015,'Length_of_Hospital_stay':.02,
    'respiratory':.015,'cardiac':.015
}.items():
    idx=rng.choice(n, size=max(1,round(n*frac)), replace=False)
    df.loc[idx,col]=np.nan
# if all PFT missing for some, Pre_op_PFTS No (rare)
all_missing = df[['FEV1','FEV_percent_OF_PREDICTED','DLCO','DLCO_percent_of_Predicted']].isna().all(axis=1)
df.loc[all_missing,'Pre_op_PFTS']='No'
raw_csv=ROOT/'data/raw/FEV1_LungCa_synthetic.csv'
df.to_csv(raw_csv,index=False)

# dictionary
labels = {
'ID':'Patient ID','Age':'Age, years','Sex':'Biological sex','Smoking':'Smoking status','Surgical_resection':'Surgical resection indicator',
'type_of_surgery':'Type of lung resection','pathologic_stage':'Pathologic stage','histology':'Tumor histology','Tumor_size_cm':'Tumor size, cm',
'Hilar_LN_involved':'Hilar lymph node involvement','Mediastinal_LN_involved':'Mediastinal lymph node involvement','Pre_op_PFTS':'Preoperative PFT available',
'FEV1':'Forced expiratory volume in 1 second, L','FEV_percent_OF_PREDICTED':'FEV1 percent predicted','DLCO':'Diffusing capacity, ml/min/mmHg',
'DLCO_percent_of_Predicted':'DLCO percent predicted','Post_op_complications':'Any postoperative complication','Length_of_Hospital_stay':'Hospital length of stay, days',
'respiratory':'Preoperative respiratory comorbidity','cardiac':'Preoperative cardiac comorbidity','stroke':'Stroke comorbidity','CKD':'Chronic kidney disease',
'DM':'Diabetes mellitus','othcomorb':'Other comorbidity','respcomp':'Respiratory postoperative complication','cardcomp':'Cardiac postoperative complication',
'deathcomp':'Postoperative death complication','infectcomp':'Infectious postoperative complication','anycomp':'Any postoperative complication',
'died':'Death indicator for overall survival','oasmons':'Overall survival/follow-up, months'}
rows=[]
for c in df.columns:
    rows.append({'Variable':c,'Label':labels.get(c,''),'Synthetic coding / units':str(df[c].dropna().iloc[0]) if c not in ['Age','Tumor_size_cm','FEV1','FEV_percent_OF_PREDICTED','DLCO','DLCO_percent_of_Predicted','Length_of_Hospital_stay','oasmons'] else labels.get(c,''),'Role':'Outcome' if c in ['died','oasmons','anycomp','respcomp','cardcomp','deathcomp','infectcomp','Length_of_Hospital_stay'] else ('Primary exposure' if c in ['FEV1','FEV_percent_OF_PREDICTED','DLCO','DLCO_percent_of_Predicted'] else 'Covariate/descriptor')})
pd.DataFrame(rows).to_csv(ROOT/'data/dictionary/synthetic_data_dictionary.csv',index=False)

# data generator script copy
shutil.copy2('/mnt/data/create_qe2022_project.py', ROOT/'scripts_generate_project.py')

# ---------------- R script ----------------
r_script = r'''# 0) Setup ----
# SEARCH KEYS: KEY_2022_PFT_SURVIVAL, KEY_FOUR_PFT_SEPARATE_MODELS, KEY_UNIVARIATE_COX
# SEARCH KEYS: KEY_PFT_COMPLICATION_LOGISTIC, KEY_LENGTH_OF_STAY_LOG_LINEAR
# SEARCH KEYS: KEY_COX_PH_SCHOENFELD, KEY_INTERACTION_EFFECT_MODIFICATION
# SEARCH KEYS: KEY_THREE_LINE_WORD_OUTPUT, KEY_MISSING_COMPLETE_CASE
# Project: 2022 PhD Qualifying Exam Applied Practice
# Data: Synthetic lung cancer resection cohort (N approximately 130)
# IMPORTANT: FEV1, FEV1 percent predicted, DLCO, and DLCO percent predicted
#            are analyzed in SEPARATE models, as required by the exam prompt.

rm(list = ls())

PROJECT_TITLE <- "QE2022 Applied Practice: Pulmonary Function, Lung Cancer Outcomes, and Survival"
DATA_PATH <- "data/raw/FEV1_LungCa_synthetic.csv"

OUT_TABLES <- "output/tables"
OUT_FIGURES <- "output/figures"
OUT_APPENDIX <- "output/appendix"
OUT_WORD <- "output/word"
for (d in c(OUT_TABLES, OUT_FIGURES, OUT_APPENDIX, OUT_WORD)) {
  dir.create(d, recursive = TRUE, showWarnings = FALSE)
}

required_packages <- c(
  "tidyverse", "janitor", "survival", "survminer", "broom",
  "flextable", "officer", "readxl", "haven", "MASS", "splines", "scales"
)
missing_packages <- required_packages[!vapply(required_packages, requireNamespace, logical(1), quietly = TRUE)]
if (length(missing_packages) > 0) {
  stop("Install missing packages first: ", paste(missing_packages, collapse = ", "), call. = FALSE)
}

# Do not attach MASS because MASS::select() masks dplyr::select().
library(tidyverse)
library(janitor)
library(survival)
library(survminer)
library(broom)
library(flextable)
library(officer)

# 0.1) Helper functions ----
fmt_p <- function(x) {
  ifelse(is.na(x), "", ifelse(x < 0.001, "<0.001", sprintf("%.3f", x)))
}
fmt_est_ci <- function(est, lo, hi, digits = 2) {
  ifelse(is.na(est), "", paste0(formatC(est, digits = digits, format = "f"), " (",
                                formatC(lo, digits = digits, format = "f"), ", ",
                                formatC(hi, digits = digits, format = "f"), ")"))
}
make_three_line <- function(ft, font_size = 8) {
  ft %>%
    flextable::theme_booktabs() %>%
    flextable::font(fontname = "Arial", part = "all") %>%
    flextable::fontsize(size = font_size, part = "all") %>%
    flextable::align(align = "center", part = "header") %>%
    flextable::align(j = 1, align = "left", part = "body") %>%
    flextable::valign(valign = "top", part = "all") %>%
    flextable::autofit()
}
add_table_to_doc <- function(doc, title, ft, note = NULL) {
  doc <- officer::body_add_par(doc, title, style = "heading 2")
  doc <- flextable::body_add_flextable(doc, value = ft)
  if (!is.null(note)) doc <- officer::body_add_par(doc, note, style = "Normal")
  officer::body_add_par(doc, "", style = "Normal")
}
write_csv_safely <- function(x, path) readr::write_csv(x, path, na = "")
safe_coxph <- function(formula, data) {
  tryCatch(survival::coxph(formula, data = data, x = TRUE, y = TRUE, model = TRUE),
           error = function(e) { message("Cox model failed: ", e$message); NULL })
}
safe_glm <- function(formula, data, family) {
  tryCatch(stats::glm(formula, data = data, family = family, model = TRUE),
           warning = function(w) { message("GLM warning: ", w$message); invokeRestart("muffleWarning") },
           error = function(e) { message("GLM failed: ", e$message); NULL })
}
num_summary <- function(data, var, label) {
  x <- data[[var]]
  tibble::tibble(
    Characteristic = label,
    N = sum(!is.na(x)),
    `Mean (SD)` = ifelse(sum(!is.na(x)) > 0, sprintf("%.1f (%.1f)", mean(x, na.rm = TRUE), sd(x, na.rm = TRUE)), ""),
    `Median (Q1, Q3)` = ifelse(sum(!is.na(x)) > 0,
      sprintf("%.1f (%.1f, %.1f)", median(x, na.rm = TRUE), quantile(x, .25, na.rm = TRUE), quantile(x, .75, na.rm = TRUE)), ""),
    Missing = sum(is.na(x))
  )
}
cat_summary <- function(data, var, label) {
  x <- data[[var]]
  denom <- sum(!is.na(x))
  tibble::tibble(Level = names(table(x, useNA = "no")), n = as.integer(table(x, useNA = "no"))) %>%
    dplyr::mutate(
      Characteristic = label,
      Summary = ifelse(denom > 0, paste0(n, " (", sprintf("%.1f", 100*n/denom), "%)"), "")
    ) %>%
    dplyr::select(Characteristic, Level, Summary)
}
tidy_cox_hr <- function(model, model_name, exposure_label = NULL) {
  if (is.null(model)) return(tibble::tibble())
  broom::tidy(model, conf.int = TRUE, exponentiate = TRUE) %>%
    dplyr::mutate(
      Model = model_name,
      `HR (95% CI)` = fmt_est_ci(estimate, conf.low, conf.high, 2),
      `P value` = fmt_p(p.value)
    ) %>%
    dplyr::select(Model, term, estimate, conf.low, conf.high, `HR (95% CI)`, `P value`)
}
tidy_or <- function(model, model_name) {
  if (is.null(model)) return(tibble::tibble())
  broom::tidy(model, conf.int = TRUE, exponentiate = TRUE) %>%
    dplyr::mutate(Model = model_name,
                  `OR (95% CI)` = fmt_est_ci(estimate, conf.low, conf.high, 2),
                  `P value` = fmt_p(p.value)) %>%
    dplyr::select(Model, term, estimate, conf.low, conf.high, `OR (95% CI)`, `P value`)
}
tidy_log_ratio <- function(model, model_name) {
  if (is.null(model)) return(tibble::tibble())
  broom::tidy(model, conf.int = TRUE) %>%
    dplyr::mutate(
      Model = model_name,
      ratio = exp(estimate), low = exp(conf.low), high = exp(conf.high),
      `Geometric mean ratio (95% CI)` = fmt_est_ci(ratio, low, high, 2),
      `P value` = fmt_p(p.value)
    ) %>%
    dplyr::select(Model, term, ratio, low, high, `Geometric mean ratio (95% CI)`, `P value`)
}

# 1) Import data ----
if (!file.exists(DATA_PATH)) stop("DATA_PATH not found: ", DATA_PATH, call. = FALSE)
if (grepl("\\.csv$", DATA_PATH, ignore.case = TRUE)) {
  dat_raw <- readr::read_csv(DATA_PATH, show_col_types = FALSE)
} else if (grepl("\\.xlsx$|\\.xls$", DATA_PATH, ignore.case = TRUE)) {
  dat_raw <- readxl::read_excel(DATA_PATH) %>% tibble::as_tibble()
} else if (grepl("\\.sas7bdat$", DATA_PATH, ignore.case = TRUE)) {
  dat_raw <- haven::read_sas(DATA_PATH) %>% tibble::as_tibble()
} else stop("Use CSV, XLS/XLSX, or SAS7BDAT.", call. = FALSE)

# 2) Clean variables and define analysis variables ----
# janitor::clean_names() converts the original mixed-case names to stable snake_case.
dat <- dat_raw %>%
  janitor::clean_names() %>%
  dplyr::mutate(dplyr::across(dplyr::where(is.character), ~ dplyr::na_if(trimws(.x), ""))) %>%
  dplyr::mutate(
    id = as.integer(id), age = as.numeric(age), tumor_size_cm = as.numeric(tumor_size_cm),
    fev1 = as.numeric(fev1), fev_percent_of_predicted = as.numeric(fev_percent_of_predicted),
    dlco = as.numeric(dlco), dlco_percent_of_predicted = as.numeric(dlco_percent_of_predicted),
    length_of_hospital_stay = as.numeric(length_of_hospital_stay),
    oasmons = as.numeric(oasmons), died = as.integer(died),
    sex = factor(sex, levels = c("Female", "Male")),
    smoking = factor(smoking, levels = c("Never", "Former", "Current")),
    type_of_surgery = factor(type_of_surgery, levels = c("Wedge/segmentectomy", "Lobectomy", "Pneumonectomy")),
    pathologic_stage = factor(pathologic_stage, levels = c("I", "II", "III"), ordered = TRUE),
    histology = factor(histology),
    dplyr::across(c(hilar_ln_involved, mediastinal_ln_involved, respiratory, cardiac,
                    stroke, ckd, dm, othcomorb, respcomp, cardcomp, deathcomp,
                    infectcomp, anycomp), ~ factor(.x, levels = c("No", "Yes")))
  ) %>%
  dplyr::mutate(
    event = dplyr::if_else(died == 1L, 1L, 0L, missing = NA_integer_),
    followup_months = oasmons,
    anycomp_bin = dplyr::case_when(anycomp == "Yes" ~ 1L, anycomp == "No" ~ 0L, TRUE ~ NA_integer_),
    respcomp_bin = dplyr::case_when(respcomp == "Yes" ~ 1L, respcomp == "No" ~ 0L, TRUE ~ NA_integer_),
    cardcomp_bin = dplyr::case_when(cardcomp == "Yes" ~ 1L, cardcomp == "No" ~ 0L, TRUE ~ NA_integer_),
    infectcomp_bin = dplyr::case_when(infectcomp == "Yes" ~ 1L, infectcomp == "No" ~ 0L, TRUE ~ NA_integer_),
    deathcomp_bin = dplyr::case_when(deathcomp == "Yes" ~ 1L, deathcomp == "No" ~ 0L, TRUE ~ NA_integer_),
    # Clinically interpretable increments for regression.
    fev1_per_0_5l = fev1 / 0.5,
    fev_pct_per_10 = fev_percent_of_predicted / 10,
    dlco_per_5 = dlco / 5,
    dlco_pct_per_10 = dlco_percent_of_predicted / 10
  )

pft_map <- tibble::tribble(
  ~raw_var, ~model_var, ~label, ~unit,
  "fev1", "fev1_per_0_5l", "FEV1", "per 0.5-L increase",
  "fev_percent_of_predicted", "fev_pct_per_10", "FEV1 percent predicted", "per 10-percentage-point increase",
  "dlco", "dlco_per_5", "DLCO", "per 5-unit increase",
  "dlco_percent_of_predicted", "dlco_pct_per_10", "DLCO percent predicted", "per 10-percentage-point increase"
)
PFT_MODEL_VARS <- pft_map$model_var
CORE_COVARIATES <- c("age", "sex", "smoking", "type_of_surgery", "pathologic_stage", "tumor_size_cm")
INTERACTION_MODIFIERS <- c("age", "sex", "smoking", "pathologic_stage", "respiratory")

# 3) Cohort flow and missingness ----
flow_counts <- tibble::tibble(
  Step = c("Raw records imported", "Positive nonmissing survival time", "Known event indicator", "Included in survival analysis"),
  N = c(nrow(dat), sum(!is.na(dat$followup_months) & dat$followup_months > 0),
        sum(!is.na(dat$event)), sum(!is.na(dat$followup_months) & dat$followup_months > 0 & !is.na(dat$event)))
)
analysis_vars <- unique(c("age","sex","smoking","type_of_surgery","pathologic_stage","histology",
                          "tumor_size_cm","hilar_ln_involved","mediastinal_ln_involved",
                          pft_map$raw_var, "length_of_hospital_stay","respiratory","cardiac","stroke","ckd","dm","othcomorb",
                          "respcomp","cardcomp","deathcomp","infectcomp","anycomp","followup_months","event"))
missingness_table <- purrr::map_dfr(analysis_vars, ~ tibble::tibble(
  Variable = .x, Missing_N = sum(is.na(dat[[.x]])), Total_N = nrow(dat),
  Missing_Percent = round(100*mean(is.na(dat[[.x]])),1)
)) %>% dplyr::arrange(dplyr::desc(Missing_Percent), Variable)
write_csv_safely(flow_counts, file.path(OUT_TABLES,"00_cohort_flow.csv"))
write_csv_safely(missingness_table, file.path(OUT_TABLES,"00_missingness.csv"))

# 4) Descriptive analysis ----
baseline_num <- dplyr::bind_rows(
  num_summary(dat,"age","Age, years"),
  num_summary(dat,"tumor_size_cm","Tumor size, cm"),
  num_summary(dat,"length_of_hospital_stay","Length of hospital stay, days")
)
baseline_cat <- dplyr::bind_rows(
  cat_summary(dat,"sex","Sex"), cat_summary(dat,"smoking","Smoking status"),
  cat_summary(dat,"type_of_surgery","Type of surgery"), cat_summary(dat,"pathologic_stage","Pathologic stage"),
  cat_summary(dat,"histology","Histology"), cat_summary(dat,"respiratory","Respiratory comorbidity"),
  cat_summary(dat,"cardiac","Cardiac comorbidity"), cat_summary(dat,"dm","Diabetes")
)
pft_summary <- purrr::map2_dfr(pft_map$raw_var, pft_map$label, ~ num_summary(dat,.x,.y))
outcome_summary <- tibble::tibble(
  Outcome = c("Deaths", "Censored", "Any postoperative complication", "Respiratory complication", "Cardiac complication", "Infectious complication", "Postoperative death complication"),
  `n (%)` = c(
    paste0(sum(dat$event==1,na.rm=TRUE)," (",sprintf("%.1f",100*mean(dat$event==1,na.rm=TRUE)),"%)"),
    paste0(sum(dat$event==0,na.rm=TRUE)," (",sprintf("%.1f",100*mean(dat$event==0,na.rm=TRUE)),"%)"),
    paste0(sum(dat$anycomp_bin==1,na.rm=TRUE)," (",sprintf("%.1f",100*mean(dat$anycomp_bin==1,na.rm=TRUE)),"%)"),
    paste0(sum(dat$respcomp_bin==1,na.rm=TRUE)," (",sprintf("%.1f",100*mean(dat$respcomp_bin==1,na.rm=TRUE)),"%)"),
    paste0(sum(dat$cardcomp_bin==1,na.rm=TRUE)," (",sprintf("%.1f",100*mean(dat$cardcomp_bin==1,na.rm=TRUE)),"%)"),
    paste0(sum(dat$infectcomp_bin==1,na.rm=TRUE)," (",sprintf("%.1f",100*mean(dat$infectcomp_bin==1,na.rm=TRUE)),"%)"),
    paste0(sum(dat$deathcomp_bin==1,na.rm=TRUE)," (",sprintf("%.1f",100*mean(dat$deathcomp_bin==1,na.rm=TRUE)),"%)")
  )
)
write_csv_safely(baseline_num, file.path(OUT_TABLES,"01_baseline_continuous.csv"))
write_csv_safely(baseline_cat, file.path(OUT_TABLES,"01_baseline_categorical.csv"))
write_csv_safely(pft_summary, file.path(OUT_TABLES,"02_pft_summary.csv"))
write_csv_safely(outcome_summary, file.path(OUT_TABLES,"03_outcome_summary.csv"))

# Distribution figures
pft_long <- dat %>% dplyr::select(dplyr::all_of(pft_map$raw_var)) %>%
  tidyr::pivot_longer(dplyr::everything(), names_to="PFT", values_to="Value")
p_pft <- ggplot2::ggplot(pft_long, ggplot2::aes(Value)) + ggplot2::geom_histogram(bins=20) +
  ggplot2::facet_wrap(~PFT, scales="free") + ggplot2::theme_bw(base_size=10) +
  ggplot2::labs(title="Distribution of preoperative pulmonary function measures", y="Count")
ggplot2::ggsave(file.path(OUT_FIGURES,"Figure_1_PFT_distributions.png"), p_pft, width=8, height=5.5, dpi=300)
p_los <- ggplot2::ggplot(dat, ggplot2::aes(length_of_hospital_stay)) + ggplot2::geom_histogram(bins=20) +
  ggplot2::theme_bw(base_size=11) + ggplot2::labs(x="Length of hospital stay, days", y="Count")
ggplot2::ggsave(file.path(OUT_APPENDIX,"Appendix_LOS_distribution.png"), p_los, width=6, height=4, dpi=300)

# 5) Overall survival: KM and univariate Cox ----
surv_dat <- dat %>% dplyr::filter(!is.na(followup_months), followup_months > 0, !is.na(event))
km_overall <- survival::survfit(survival::Surv(followup_months,event) ~ 1, data=surv_dat)
km_plot <- survminer::ggsurvplot(km_overall, data=surv_dat, conf.int=TRUE, risk.table=TRUE,
  xlab="Months after surgery", ylab="Overall survival probability", break.time.by=12,
  ggtheme=ggplot2::theme_bw(base_size=11), risk.table.height=.25)
ggplot2::ggsave(file.path(OUT_FIGURES,"Figure_2_Overall_Kaplan_Meier.png"), km_plot$plot, width=7, height=5, dpi=300)
# Save combined KM + risk table safely
png(file.path(OUT_FIGURES,"Figure_2_Overall_Kaplan_Meier_with_risk_table.png"), width=2100,height=1800,res=300)
print(km_plot)
dev.off()

univ_vars <- c("age","sex","smoking","type_of_surgery","pathologic_stage","histology","tumor_size_cm",
               "hilar_ln_involved","mediastinal_ln_involved","respiratory","cardiac","stroke","ckd","dm","othcomorb", PFT_MODEL_VARS)
univ_results <- purrr::map_dfr(univ_vars, function(v) {
  d <- dat %>% dplyr::select(followup_months,event,dplyr::all_of(v)) %>% tidyr::drop_na()
  m <- safe_coxph(stats::as.formula(paste0("Surv(followup_months,event) ~ `",v,"`")), d)
  tidy_cox_hr(m, paste0("Univariate: ",v))
})
write_csv_safely(univ_results, file.path(OUT_TABLES,"04_univariate_cox.csv"))

# 6) Four separate adjusted Cox models ----
fit_adjusted_cox <- function(model_var, label, unit) {
  needed <- c("followup_months","event",model_var,CORE_COVARIATES)
  d <- dat %>% dplyr::select(dplyr::all_of(needed)) %>% tidyr::drop_na()
  f <- stats::as.formula(paste("Surv(followup_months,event) ~", paste(c(model_var,CORE_COVARIATES),collapse=" + ")))
  m <- safe_coxph(f,d)
  list(data=d, model=m, result=tidy_cox_hr(m,paste0(label," ",unit)))
}
cox_fits <- purrr::pmap(pft_map[,c("model_var","label","unit")], fit_adjusted_cox)
names(cox_fits) <- pft_map$model_var
adjusted_cox_results <- purrr::map_dfr(cox_fits,"result")
write_csv_safely(adjusted_cox_results, file.path(OUT_TABLES,"05_adjusted_cox_four_separate_pft_models.csv"))

# Compact primary-exposure rows for the report
primary_cox_rows <- purrr::map2_dfr(pft_map$model_var, seq_len(nrow(pft_map)), function(v,i) {
  x <- cox_fits[[v]]$result %>% dplyr::filter(term == v)
  if (nrow(x)==0) return(tibble::tibble())
  tibble::tibble(PFT=pft_map$label[i], Increment=pft_map$unit[i], N=nrow(cox_fits[[v]]$data),
                 `Adjusted HR (95% CI)`=x$`HR (95% CI)`, `P value`=x$`P value`)
})
write_csv_safely(primary_cox_rows, file.path(OUT_TABLES,"05_primary_pft_adjusted_hr.csv"))

# Forest plot of four primary PFT effects
forest_dat <- purrr::map2_dfr(pft_map$model_var, seq_len(nrow(pft_map)), function(v,i) {
  x <- cox_fits[[v]]$result %>% dplyr::filter(term==v)
  if (nrow(x)==0) return(tibble::tibble())
  tibble::tibble(PFT=paste0(pft_map$label[i]," (",pft_map$unit[i],")"), HR=x$estimate, low=x$conf.low, high=x$conf.high)
})
if (nrow(forest_dat)>0) {
  p_forest <- ggplot2::ggplot(forest_dat, ggplot2::aes(x=HR,y=reorder(PFT,HR))) +
    ggplot2::geom_vline(xintercept=1,linetype=2) + ggplot2::geom_point() +
    ggplot2::geom_errorbarh(ggplot2::aes(xmin=low,xmax=high),height=.2) +
    ggplot2::scale_x_log10() + ggplot2::theme_bw(base_size=11) +
    ggplot2::labs(x="Adjusted hazard ratio (log scale)",y=NULL,title="Adjusted associations of pulmonary function with overall survival")
  ggplot2::ggsave(file.path(OUT_FIGURES,"Figure_3_Adjusted_PFT_HR_Forest.png"),p_forest,width=7,height=4.5,dpi=300)
}

# 7) Cox PH assumptions and continuous functional form ----
ph_results <- purrr::imap_dfr(cox_fits, function(obj,name) {
  if (is.null(obj$model)) return(tibble::tibble())
  z <- survival::cox.zph(obj$model)
  as.data.frame(z$table) %>% tibble::rownames_to_column("Term") %>% tibble::as_tibble() %>%
    dplyr::transmute(Model=name,Term,Chisq=round(chisq,3),df=round(df,0),`P value`=fmt_p(p))
})
write_csv_safely(ph_results,file.path(OUT_TABLES,"06_cox_ph_assumption_tests.csv"))
# Save one PH plot file per model
purrr::iwalk(cox_fits,function(obj,name){
  if (!is.null(obj$model)) {
    z <- survival::cox.zph(obj$model)
    grDevices::png(file.path(OUT_APPENDIX,paste0("Appendix_PH_",name,".png")),width=2400,height=1800,res=250)
    plot(z); grDevices::dev.off()
  }
})

# 8) Interaction/effect-modification screening ----
fit_interaction <- function(model_var, modifier, pft_label) {
  needed <- unique(c("followup_months","event",model_var,modifier,CORE_COVARIATES))
  d <- dat %>% dplyr::select(dplyr::all_of(needed)) %>% tidyr::drop_na()
  adjustment <- setdiff(CORE_COVARIATES, modifier)
  f0 <- stats::as.formula(paste("Surv(followup_months,event) ~",paste(c(model_var,modifier,adjustment),collapse=" + ")))
  f1 <- stats::as.formula(paste("Surv(followup_months,event) ~",paste(c(paste0(model_var," * ",modifier),adjustment),collapse=" + ")))
  m0 <- safe_coxph(f0,d); m1 <- safe_coxph(f1,d)
  if (is.null(m0)||is.null(m1)) return(tibble::tibble(PFT=pft_label,Modifier=modifier,N=nrow(d),LRT=NA_real_,df=NA_real_,P_value=NA_real_,`P value`="Model failed"))
  a <- stats::anova(m0,m1,test="LRT")
  p <- as.numeric(a$`Pr(>|Chi|)`[2])
  tibble::tibble(PFT=pft_label,Modifier=modifier,N=nrow(d),LRT=round(as.numeric(a$Chisq[2]),3),df=as.numeric(a$Df[2]),P_value=p,`P value`=fmt_p(p))
}
interaction_results <- purrr::map_dfr(seq_len(nrow(pft_map)), function(i) {
  purrr::map_dfr(INTERACTION_MODIFIERS, ~ fit_interaction(pft_map$model_var[i],.x,pft_map$label[i]))
}) %>% dplyr::arrange(P_value)
write_csv_safely(interaction_results,file.path(OUT_TABLES,"07_pft_interaction_screening.csv"))

# 9) PFT and postoperative complications ----
comp_outcomes <- c(anycomp_bin="Any postoperative complication",respcomp_bin="Respiratory complication",
                   cardcomp_bin="Cardiac complication",infectcomp_bin="Infectious complication",deathcomp_bin="Postoperative death complication")
# Unadjusted comparisons: median PFT by outcome + Wilcoxon test
pft_comp_screen <- purrr::map_dfr(names(comp_outcomes), function(outcome) {
  purrr::map_dfr(seq_len(nrow(pft_map)), function(i) {
    raw <- pft_map$raw_var[i]
    d <- dat %>% dplyr::select(dplyr::all_of(c(outcome,raw))) %>% tidyr::drop_na()
    if (length(unique(d[[outcome]]))<2) return(tibble::tibble())
    wt <- stats::wilcox.test(d[[raw]] ~ d[[outcome]], exact=FALSE)
    tibble::tibble(Outcome=comp_outcomes[[outcome]],PFT=pft_map$label[i],N=nrow(d),
      `No complication, median (IQR)`=sprintf("%.1f (%.1f, %.1f)",median(d[[raw]][d[[outcome]]==0]),quantile(d[[raw]][d[[outcome]]==0],.25),quantile(d[[raw]][d[[outcome]]==0],.75)),
      `Complication, median (IQR)`=sprintf("%.1f (%.1f, %.1f)",median(d[[raw]][d[[outcome]]==1]),quantile(d[[raw]][d[[outcome]]==1],.25),quantile(d[[raw]][d[[outcome]]==1],.75)),
      P_value=wt$p.value,`P value`=fmt_p(wt$p.value))
  })
})
write_csv_safely(pft_comp_screen,file.path(OUT_TABLES,"08_pft_complication_unadjusted_screening.csv"))

# Adjusted logistic models for ANY complication, one PFT at a time.
fit_adjusted_logistic <- function(model_var,label,unit) {
  needed <- c("anycomp_bin",model_var,CORE_COVARIATES)
  d <- dat %>% dplyr::select(dplyr::all_of(needed)) %>% tidyr::drop_na()
  f <- stats::as.formula(paste("anycomp_bin ~",paste(c(model_var,CORE_COVARIATES),collapse=" + ")))
  m <- safe_glm(f,d,stats::binomial())
  list(data=d,model=m,result=tidy_or(m,paste0(label," ",unit)))
}
logit_fits <- purrr::pmap(pft_map[,c("model_var","label","unit")],fit_adjusted_logistic)
names(logit_fits)<-pft_map$model_var
logit_primary <- purrr::map2_dfr(pft_map$model_var,seq_len(nrow(pft_map)),function(v,i){
  x<-logit_fits[[v]]$result %>% dplyr::filter(term==v)
  if(nrow(x)==0)return(tibble::tibble())
  tibble::tibble(PFT=pft_map$label[i],Increment=pft_map$unit[i],N=nrow(logit_fits[[v]]$data),`Adjusted OR (95% CI)`=x$`OR (95% CI)`,`P value`=x$`P value`)
})
write_csv_safely(logit_primary,file.path(OUT_TABLES,"09_adjusted_logistic_any_complication.csv"))

# Optional selected secondary complication models when unadjusted P < 0.10 and >=10 events.
selected_secondary <- pft_comp_screen %>% dplyr::filter(Outcome!="Any postoperative complication",P_value<.10)
secondary_logit_results <- purrr::pmap_dfr(selected_secondary[,c("Outcome","PFT")], function(Outcome,PFT) {
  outcome <- names(comp_outcomes)[match(Outcome,unname(comp_outcomes))]
  i <- match(PFT,pft_map$label); mv<-pft_map$model_var[i]
  needed<-c(outcome,mv,CORE_COVARIATES); d<-dat %>% dplyr::select(dplyr::all_of(needed)) %>% tidyr::drop_na()
  if(sum(d[[outcome]]==1)<10)return(tibble::tibble(Outcome=Outcome,PFT=PFT,Note="Fewer than 10 events; adjusted model not fit."))
  f<-stats::as.formula(paste(outcome,"~",paste(c(mv,CORE_COVARIATES),collapse=" + ")))
  m<-safe_glm(f,d,stats::binomial()); x<-tidy_or(m,paste(Outcome,PFT)) %>% dplyr::filter(term==mv)
  if(nrow(x)==0)return(tibble::tibble(Outcome=Outcome,PFT=PFT,Note="Model failed."))
  tibble::tibble(Outcome=Outcome,PFT=PFT,N=nrow(d),`Adjusted OR (95% CI)`=x$`OR (95% CI)`,`P value`=x$`P value`,Note="")
})
write_csv_safely(secondary_logit_results,file.path(OUT_TABLES,"10_selected_secondary_complication_models.csv"))

# 10) Length of stay ----
# LOS is right-skewed in most clinical datasets. Keep it continuous and model log(LOS)
# unless there is a prespecified clinically meaningful categorization.
fit_los_model <- function(model_var,label,unit) {
  needed<-c("length_of_hospital_stay",model_var,CORE_COVARIATES)
  d<-dat %>% dplyr::select(dplyr::all_of(needed)) %>% tidyr::drop_na() %>% dplyr::filter(length_of_hospital_stay>0)
  f<-stats::as.formula(paste("log(length_of_hospital_stay) ~",paste(c(model_var,CORE_COVARIATES),collapse=" + ")))
  m<-stats::lm(f,d)
  list(data=d,model=m,result=tidy_log_ratio(m,paste0(label," ",unit)))
}
los_fits<-purrr::pmap(pft_map[,c("model_var","label","unit")],fit_los_model); names(los_fits)<-pft_map$model_var
los_primary<-purrr::map2_dfr(pft_map$model_var,seq_len(nrow(pft_map)),function(v,i){
  x<-los_fits[[v]]$result %>% dplyr::filter(term==v)
  if(nrow(x)==0)return(tibble::tibble())
  tibble::tibble(PFT=pft_map$label[i],Increment=pft_map$unit[i],N=nrow(los_fits[[v]]$data),
                 `Adjusted geometric mean ratio (95% CI)`=x$`Geometric mean ratio (95% CI)`,`P value`=x$`P value`)
})
write_csv_safely(los_primary,file.path(OUT_TABLES,"11_adjusted_los_log_linear_models.csv"))
# Gamma log-link sensitivity
los_gamma<-purrr::map2_dfr(pft_map$model_var,seq_len(nrow(pft_map)),function(v,i){
  needed<-c("length_of_hospital_stay",v,CORE_COVARIATES); d<-dat %>% dplyr::select(dplyr::all_of(needed)) %>% tidyr::drop_na() %>% dplyr::filter(length_of_hospital_stay>0)
  f<-stats::as.formula(paste("length_of_hospital_stay ~",paste(c(v,CORE_COVARIATES),collapse=" + ")))
  m<-safe_glm(f,d,stats::Gamma(link="log")); x<-tidy_or(m,paste0(pft_map$label[i]," Gamma sensitivity")) %>% dplyr::filter(term==v)
  if(nrow(x)==0)return(tibble::tibble())
  tibble::tibble(PFT=pft_map$label[i],Increment=pft_map$unit[i],N=nrow(d),`Adjusted mean ratio (95% CI)`=x$`OR (95% CI)`,`P value`=x$`P value`)
})
write_csv_safely(los_gamma,file.path(OUT_TABLES,"12_los_gamma_sensitivity.csv"))

# LOS residual diagnostics for first PFT model (repeat as needed)
first_los <- los_fits[[1]]$model
if (!is.null(first_los)) {
  png(file.path(OUT_APPENDIX,"Appendix_LOS_residual_diagnostics.png"),width=2200,height=1800,res=250)
  par(mfrow=c(2,2)); plot(first_los); dev.off()
}

# 11) Create three-line Word documents ----
ft_baseline_num <- make_three_line(flextable::flextable(baseline_num),8)
ft_baseline_cat <- make_three_line(flextable::flextable(baseline_cat),8)
ft_pft <- make_three_line(flextable::flextable(pft_summary),8)
ft_outcomes <- make_three_line(flextable::flextable(outcome_summary),8)
ft_univ_cox <- make_three_line(flextable::flextable(univ_results %>% dplyr::select(Model,term,`HR (95% CI)`,`P value`)),7)
ft_primary_cox <- make_three_line(flextable::flextable(primary_cox_rows),8)
ft_interactions <- make_three_line(flextable::flextable(interaction_results %>% dplyr::select(PFT,Modifier,N,LRT,df,`P value`)),7)
ft_comp_screen <- make_three_line(flextable::flextable(pft_comp_screen %>% dplyr::select(Outcome,PFT,N,`No complication, median (IQR)`,`Complication, median (IQR)`,`P value`)),7)
ft_logit <- make_three_line(flextable::flextable(logit_primary),8)
ft_los <- make_three_line(flextable::flextable(los_primary),8)

main_doc <- officer::read_docx()
main_doc <- officer::body_add_par(main_doc,PROJECT_TITLE,style="heading 1")
main_doc <- officer::body_add_par(main_doc,paste0("Generated: ",Sys.Date()),style="Normal")
main_doc <- officer::body_add_par(main_doc,"Synthetic practice data only. Interpretations must be written from the actual final output.",style="Normal")
main_doc <- add_table_to_doc(main_doc,"Table 1A. Baseline continuous characteristics",ft_baseline_num)
main_doc <- add_table_to_doc(main_doc,"Table 1B. Baseline categorical characteristics",ft_baseline_cat)
main_doc <- add_table_to_doc(main_doc,"Table 2. Pulmonary function test summary",ft_pft,note="The four PFT measures should not be entered together in the same multivariable model.")
main_doc <- add_table_to_doc(main_doc,"Table 3. Outcome summary",ft_outcomes)
for (img_title in c("Figure 1. Pulmonary function distributions"="Figure_1_PFT_distributions.png",
                    "Figure 2. Overall Kaplan-Meier survival curve"="Figure_2_Overall_Kaplan_Meier_with_risk_table.png",
                    "Figure 3. Adjusted PFT hazard-ratio forest plot"="Figure_3_Adjusted_PFT_HR_Forest.png")) {
  path<-file.path(OUT_FIGURES,img_title)
}
# Add figures explicitly
if(file.exists(file.path(OUT_FIGURES,"Figure_1_PFT_distributions.png"))){main_doc<-officer::body_add_par(main_doc,"Figure 1. Pulmonary function distributions",style="heading 2");main_doc<-officer::body_add_img(main_doc,file.path(OUT_FIGURES,"Figure_1_PFT_distributions.png"),width=6.8,height=4.7)}
if(file.exists(file.path(OUT_FIGURES,"Figure_2_Overall_Kaplan_Meier_with_risk_table.png"))){main_doc<-officer::body_add_par(main_doc,"Figure 2. Overall Kaplan-Meier survival curve",style="heading 2");main_doc<-officer::body_add_img(main_doc,file.path(OUT_FIGURES,"Figure_2_Overall_Kaplan_Meier_with_risk_table.png"),width=6.8,height=5.4)}
if(file.exists(file.path(OUT_FIGURES,"Figure_3_Adjusted_PFT_HR_Forest.png"))){main_doc<-officer::body_add_par(main_doc,"Figure 3. Adjusted associations of PFT with overall survival",style="heading 2");main_doc<-officer::body_add_img(main_doc,file.path(OUT_FIGURES,"Figure_3_Adjusted_PFT_HR_Forest.png"),width=6.7,height=4.3)}
main_doc <- add_table_to_doc(main_doc,"Table 4. Univariate Cox proportional hazards models",ft_univ_cox)
main_doc <- add_table_to_doc(main_doc,"Table 5. Four separate adjusted Cox models for pulmonary function",ft_primary_cox,note="Each row is from a separate adjusted Cox model. HRs are expressed per clinically interpretable PFT increment.")
main_doc <- add_table_to_doc(main_doc,"Table 6. PFT effect-modification screening",ft_interactions,note="Likelihood-ratio tests compare models with and without the stated PFT-by-modifier interaction.")
main_doc <- add_table_to_doc(main_doc,"Table 7. Unadjusted PFT comparisons by postoperative complication",ft_comp_screen)
main_doc <- add_table_to_doc(main_doc,"Table 8. Adjusted logistic models for any postoperative complication",ft_logit,note="Each row is from a separate multivariable logistic model; OR must not be described as a risk ratio.")
main_doc <- add_table_to_doc(main_doc,"Table 9. Adjusted log-linear models for length of stay",ft_los,note="Exponentiated coefficients are ratios of geometric mean length of stay.")
print(main_doc,target=file.path(OUT_WORD,"QE2022_Lung_Cancer_PFT_Main_Tables_ThreeLine.docx"))

# Appendix
ft_flow <- make_three_line(flextable::flextable(flow_counts),9)
ft_missing <- make_three_line(flextable::flextable(missingness_table),8)
ft_ph <- make_three_line(flextable::flextable(ph_results),7)
ft_secondary <- make_three_line(flextable::flextable(secondary_logit_results),7)
ft_gamma <- make_three_line(flextable::flextable(los_gamma),8)
appendix_doc <- officer::read_docx()
appendix_doc <- officer::body_add_par(appendix_doc,"Appendix: Diagnostics and Sensitivity Analyses",style="heading 1")
appendix_doc <- add_table_to_doc(appendix_doc,"Appendix Table A1. Cohort flow",ft_flow)
appendix_doc <- add_table_to_doc(appendix_doc,"Appendix Table A2. Missingness summary",ft_missing)
appendix_doc <- add_table_to_doc(appendix_doc,"Appendix Table A3. Cox proportional hazards assumption tests",ft_ph)
appendix_doc <- add_table_to_doc(appendix_doc,"Appendix Table A4. Selected secondary complication models",ft_secondary)
appendix_doc <- add_table_to_doc(appendix_doc,"Appendix Table A5. Gamma log-link LOS sensitivity models",ft_gamma)
for(img in list.files(OUT_APPENDIX,pattern="\\.png$",full.names=TRUE)){
  appendix_doc<-officer::body_add_par(appendix_doc,tools::file_path_sans_ext(basename(img)),style="heading 2")
  appendix_doc<-officer::body_add_img(appendix_doc,img,width=6.6,height=4.8)
}
print(appendix_doc,target=file.path(OUT_WORD,"QE2022_Lung_Cancer_PFT_Appendix_Diagnostics_ThreeLine.docx"))

# 12) Session information ----
writeLines(capture.output(sessionInfo()), file.path(OUT_APPENDIX,"sessionInfo.txt"))
message("Analysis complete. Review output/word, output/tables, output/figures, and output/appendix.")
'''
(ROOT/'scripts/QE2022_Applied_Lung_Cancer_PFT_full_analysis.R').write_text(r_script, encoding='utf-8')

install_r = r'''# 00) Install required packages ----
packages <- c("tidyverse","janitor","survival","survminer","broom","flextable","officer","readxl","haven","MASS","splines","scales")
missing <- packages[!vapply(packages, requireNamespace, logical(1), quietly=TRUE)]
if(length(missing)>0) install.packages(missing)
message("Package check complete.")
'''
(ROOT/'scripts/00_install_required_packages.R').write_text(install_r, encoding='utf-8')

# ---------------- README ----------------
readme = f'''# QE2022 Applied: Lung Cancer PFT Project

This is a **synthetic practice project** based on the structure of the 2022 Applied Qualifying Exam. It is not the original clinical dataset and must not be interpreted as real clinical evidence.

## Open and run

1. Unzip the project.
2. Open `QE2022_Applied_Lung_Cancer_PFT_Project.Rproj` in RStudio.
3. If needed, run `source("scripts/00_install_required_packages.R")` once.
4. Run:

```r
source("scripts/QE2022_Applied_Lung_Cancer_PFT_full_analysis.R")
```

The data path is already set to:

```r
DATA_PATH <- "data/raw/FEV1_LungCa_synthetic.csv"
```

## Project contents

- `data/raw/`: synthetic lung cancer dataset, N={n}
- `data/dictionary/`: synthetic data dictionary
- `scripts/`: complete end-to-end analysis and package installer
- `docs/`: Chinese guide, blank report template, original exam, and general reference handbook
- `output/`: tables, figures, appendix diagnostics, and Word three-line tables

## Key analytic decisions

- Overall survival: Kaplan-Meier plus Cox proportional hazards models.
- The four PFT measures are analyzed in **four separate models**, not simultaneously.
- Continuous PFT effects are reported per clinically interpretable increments.
- Postoperative complications: unadjusted PFT comparisons and separate adjusted logistic models.
- Length of stay: distribution check plus log-linear model; Gamma log-link sensitivity analysis.
- Interaction screening: one PFT-by-modifier interaction at a time using likelihood-ratio tests.
- Missing data: complete-case analysis within each model, with model-specific N reported.

## Search keys

Search the R script for:

- `KEY_2022_PFT_SURVIVAL`
- `KEY_FOUR_PFT_SEPARATE_MODELS`
- `KEY_PFT_COMPLICATION_LOGISTIC`
- `KEY_LENGTH_OF_STAY_LOG_LINEAR`
- `KEY_COX_PH_SCHOENFELD`
- `KEY_INTERACTION_EFFECT_MODIFICATION`
- `KEY_THREE_LINE_WORD_OUTPUT`

## Important exam reminder

The script is a learning and navigation tool. During the exam, you must independently verify variable definitions, model assumptions, effect measures, reference groups, and written interpretations.
'''
(ROOT/'README.md').write_text(readme, encoding='utf-8')
(ROOT/'QE2022_Applied_Lung_Cancer_PFT_Project.Rproj').write_text('Version: 1.0\n\nRestoreWorkspace: No\nSaveWorkspace: No\nAlwaysSaveHistory: No\n\nEnableCodeIndexing: Yes\nUseSpacesForTab: Yes\nNumSpacesForTab: 2\nEncoding: UTF-8\n\nRnwWeave: Sweave\nLaTeX: pdfLaTeX\n', encoding='utf-8')
(ROOT/'.gitignore').write_text('.Rhistory\n.RData\n.Ruserdata\n.Rproj.user/\noutput/tables/*\noutput/figures/*\noutput/appendix/*\noutput/word/*\n!output/**/.gitkeep\n', encoding='utf-8')
for p in [ROOT/'output/tables/.gitkeep',ROOT/'output/figures/.gitkeep',ROOT/'output/appendix/.gitkeep',ROOT/'output/word/.gitkeep']:
    p.write_text('',encoding='utf-8')

# Copy source and general references
shutil.copy2('/mnt/data/QE2022-PhD-Applied.docx', ROOT/'docs/source/QE2022-PhD-Applied.docx')
for src in ['/mnt/data/Applied_Biostatistics_Research_to_Report_Handbook_CN.docx','/mnt/data/Applied_Biostatistics_Research_to_Report_Quick_Reference.xlsx']:
    if Path(src).exists(): shutil.copy2(src, ROOT/'docs/general_reference'/Path(src).name)

# ---------------- DOCX helpers ----------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_border(cell, **kwargs):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            edge_data=kwargs.get(edge); tag='w:{}'.format(edge); element=tcBorders.find(qn(tag))
            if element is None: element=OxmlElement(tag); tcBorders.append(element)
            for key,val in edge_data.items(): element.set(qn('w:'+key),str(val))

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def style_three_line_table(table, header_fill='D9EAF7'):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx,row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for run in p.runs: run.font.name='Arial'; run.font.size=Pt(8.5)
            # remove all borders
            set_cell_border(cell,top={'val':'nil'},bottom={'val':'nil'},left={'val':'nil'},right={'val':'nil'})
        if r_idx==0:
            set_repeat_table_header(row)
            for cell in row.cells:
                set_cell_shading(cell,header_fill)
                set_cell_border(cell,top={'val':'single','sz':'10','color':'000000'},bottom={'val':'single','sz':'8','color':'000000'})
                for p in cell.paragraphs:
                    for run in p.runs: run.bold=True
    for cell in table.rows[-1].cells:
        set_cell_border(cell,bottom={'val':'single','sz':'10','color':'000000'})

def add_code(doc, code):
    p=doc.add_paragraph()
    p.style=doc.styles['No Spacing']
    p.paragraph_format.left_indent=Inches(.18)
    p.paragraph_format.right_indent=Inches(.12)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(5)
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F3F6F8'); pPr.append(shd)
    r=p.add_run(code); r.font.name='Consolas'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(30,30,30)

def add_tip(doc,title,text,fill='E2F0D9'):
    table=doc.add_table(rows=1,cols=1); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=table.cell(0,0); set_cell_shading(cell,fill); set_cell_border(cell,top={'val':'single','sz':'6','color':'6B8E23'},bottom={'val':'single','sz':'6','color':'6B8E23'},left={'val':'single','sz':'6','color':'6B8E23'},right={'val':'single','sz':'6','color':'6B8E23'})
    p=cell.paragraphs[0]; r=p.add_run(title+'：'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(9)
    r2=p.add_run(text); r2.font.name='Arial'; r2.font.size=Pt(9)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def setup_doc(doc,title,subtitle):
    sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.7); sec.right_margin=Inches(.7)
    styles=doc.styles
    styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10.5)
    for sty,size,color in [('Title',22,'17365D'),('Heading 1',16,'17365D'),('Heading 2',13,'1F6D70'),('Heading 3',11,'1F6D70')]:
        styles[sty].font.name='Arial'; styles[sty].font.size=Pt(size); styles[sty].font.color.rgb=RGBColor.from_string(color); styles[sty].font.bold=True
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.name='Arial'; r.font.size=Pt(22); r.font.color.rgb=RGBColor(23,54,93)
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p2.add_run(subtitle); r.font.name='Arial'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(31,109,112)
    doc.add_paragraph('Synthetic practice project | 中文讲解 + English report templates',style='Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def add_page_number(section):
    footer=section.footer
    p=footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run('Page ')
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text='PAGE'
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

# Guide DOCX
guide=Document(); setup_doc(guide,'QE2022 Applied：肺癌 PFT 项目中文指南','研究问题 → 数据准备 → Cox / Logistic / LOS → Diagnostics → Report')
for sec in guide.sections: add_page_number(sec)

guide.add_heading('使用说明',level=1)
guide.add_paragraph('这份指南对应 2022 Applied 题型。目标不是背代码，而是理解为什么每个 block 存在、输出看哪里、何时需要换方法，以及如何把结果写成正式 technical report。')
add_tip(guide,'最重要规则','FEV1、FEV1 percent predicted、DLCO、DLCO percent predicted 高度相关且代表不同尺度，题目明确要求四者不要进入同一个模型。正确做法是建立四个平行、结构相同的模型。')

guide.add_heading('目录与项目导航',level=1)
nav=[('data/raw','synthetic practice data'),('scripts','完整端到端 R script'),('output/tables','CSV tables'),('output/figures','主文图'),('output/appendix','diagnostic plots and sessionInfo'),('output/word','三线表 Word output'),('docs/general_reference','通用 Applied 手册与 Excel 导航表')]
t=guide.add_table(rows=1,cols=2); t.rows[0].cells[0].text='Folder'; t.rows[0].cells[1].text='Purpose'
for a,b in nav: row=t.add_row().cells; row[0].text=a; row[1].text=b
style_three_line_table(t)

guide.add_heading('1. 先把题目翻译成变量地图',level=1)
for s in [
'Population：接受肺癌手术切除的患者。',
'Primary exposures：四个连续 PFT 指标；每个分别建模。',
'Primary outcome：overall survival，包含 follow-up time 和 death indicator，存在右删失。',
'Secondary outcomes：respiratory、cardiac、infectious、death、any postoperative complication；均为 binary。',
'Tertiary outcome：length of hospital stay，通常右偏且保持连续比随意二分类更好。',
'Covariates：age、sex、smoking、type of surgery、stage、histology、tumor size、comorbidities。',
'Effect modifiers：题目问“effect differs by”，对应 PFT × modifier interaction。']:
    guide.add_paragraph(s,style='List Bullet')

guide.add_heading('2. Block 0：Setup、packages 与输出目录',level=1)
guide.add_paragraph('这一块负责保证脚本可以从干净 R session 运行，并把所有结果放进固定文件夹。考试时最忌讳把对象散落在 Global Environment，或依赖前一天残留对象。')
add_code(guide,'rm(list = ls())\nDATA_PATH <- "data/raw/FEV1_LungCa_synthetic.csv"\nfor (d in c(OUT_TABLES, OUT_FIGURES, OUT_APPENDIX, OUT_WORD)) {\n  dir.create(d, recursive = TRUE, showWarnings = FALSE)\n}')
add_tip(guide,'函数冲突','不要 attach MASS。MASS::select() 会遮蔽 dplyr::select()。脚本对容易冲突的函数显式写 dplyr::select()、dplyr::filter() 和 dplyr::mutate()。','FFF2CC')

guide.add_heading('3. Block 1-2：导入、clean_names 和变量编码',level=1)
guide.add_paragraph('导入后先用 janitor::clean_names() 统一变量名，再明确 factor level 和 reference group。reference group 决定 OR、HR 和 regression coefficient 的比较方向。')
add_code(guide,'dat <- dat_raw %>%\n  janitor::clean_names() %>%\n  dplyr::mutate(\n    sex = factor(sex, levels = c("Female", "Male")),\n    smoking = factor(smoking, levels = c("Never", "Former", "Current")),\n    pathologic_stage = factor(pathologic_stage, levels = c("I", "II", "III"), ordered = TRUE),\n    event = dplyr::if_else(died == 1L, 1L, 0L),\n    followup_months = oasmons\n  )')
add_tip(guide,'检查顺序','先运行 names(dat_raw)、str(dat_raw)、table(..., useNA="ifany")；确认原始编码后再 factor。不要根据想当然设置 reference。')

guide.add_heading('4. 为什么要把 PFT 缩放',level=1)
guide.add_paragraph('如果直接报告 DLCO percent predicted 每增加 1 个百分点的 HR，效应往往接近 1，不利于解释。脚本把 FEV1 按 0.5 L、百分比指标按 10 percentage points、DLCO 按 5 units 缩放。缩放只改变解释单位，不改变拟合或 P value。')
add_code(guide,'fev1_per_0_5l = fev1 / 0.5\nfev_pct_per_10 = fev_percent_of_predicted / 10\ndlco_per_5 = dlco / 5\ndlco_pct_per_10 = dlco_percent_of_predicted / 10')

guide.add_heading('5. Cohort flow、missingness 与 complete-case N',level=1)
guide.add_paragraph('每个模型可能因暴露变量缺失而使用不同 N。必须报告 model-specific N，不能只写原始样本量 130。complete-case analysis 是可执行的默认方案，但要在 limitations 中说明 missing-at-random 并未得到保证。')
add_tip(guide,'何时考虑 multiple imputation','当关键变量缺失比例不低、complete-case 会明显损失样本，且缺失机制可由已观测变量合理预测时考虑。生存结局的 MI 还应纳入 event 和 cumulative hazard 等信息；考试时间有限时至少说明。','D9EAF7')

guide.add_heading('6. Descriptive tables 和 figures',level=1)
guide.add_paragraph('题目要求三类 descriptive outputs：研究人群、PFT、outcomes。连续变量同时输出 mean/SD 与 median/IQR，能让你判断分布；正式正文可根据分布选择主要 summary。')
rows=[('Table 1','Age, sex, smoking, surgery, stage, histology, tumor size, comorbidities'),('Table 2','Four PFT distributions and missingness'),('Table 3','Events/censoring, complications, LOS'),('Figure 1','PFT histograms'),('Figure 2','Overall Kaplan-Meier curve')]
t=guide.add_table(rows=1,cols=2); t.rows[0].cells[0].text='Output'; t.rows[0].cells[1].text='Purpose'
for a,b in rows: r=t.add_row().cells; r[0].text=a; r[1].text=b
style_three_line_table(t)

guide.add_heading('7. Survival：KM、univariate Cox、adjusted Cox',level=1)
guide.add_heading('7.1 KM 回答什么',level=2)
guide.add_paragraph('Overall KM 描述整个队列的 survival experience；它不调整 confounding。若把连续 PFT 人为分成高低组画 KM，只能用于展示，不能替代连续 Cox 模型。')
guide.add_heading('7.2 四个 adjusted Cox 模型',level=2)
add_code(guide,'Surv(followup_months, event) ~ PFT + age + sex + smoking +\n  type_of_surgery + pathologic_stage + tumor_size_cm')
guide.add_paragraph('四个模型只替换 PFT，其他 covariates 保持一致，便于比较。主结果提取 PFT 那一行的 adjusted HR、95% CI 和 P value。')
add_tip(guide,'HR 的正确解释','HR 是任一时点瞬时 hazard 的相对值，不是 risk ratio，也不是 survival probability difference。例如 HR=0.80 表示在 PH 假设下，较高 PFT 与约 20% 较低的死亡 hazard 相关。')

guide.add_heading('7.3 R 输出看哪里',level=2)
for s in ['term：变量或 factor level；先确认 reference。','estimate（exponentiated）：HR。','conf.low/conf.high：95% CI；若跨 1，数据与无效应相容。','p.value：检验 coefficient=0 或 HR=1；不能替代 estimate 和 CI。','n：实际进入该模型的 complete cases。']:
    guide.add_paragraph(s,style='List Bullet')

guide.add_heading('7.4 Methods 与 Results 英文模板',level=2)
guide.add_paragraph('Methods template:')
add_code(guide,'Overall survival was summarized using the Kaplan-Meier method. Associations between each pulmonary function measure and overall survival were evaluated in separate Cox proportional hazards models adjusted for age, sex, smoking status, type of surgery, pathologic stage, and tumor size. Hazard ratios (HRs) and 95% confidence intervals (CIs) were reported per prespecified clinically interpretable increment of each pulmonary function measure.')
guide.add_paragraph('Results template:')
add_code(guide,'After multivariable adjustment, a 10-percentage-point increase in DLCO percent predicted was associated with [lower/higher] mortality hazard (adjusted HR, X.XX; 95% CI, X.XX-X.XX; P=X.XXX).')

guide.add_heading('8. Cox diagnostics：PH 和 functional form',level=1)
guide.add_paragraph('cox.zph() 检查 coefficient 是否随时间系统变化。先看 GLOBAL，再看 individual terms，并结合 Schoenfeld residual plots。P<0.05 是提示，不应机械宣布模型失败。')
add_code(guide,'zph <- survival::cox.zph(cox_model)\nprint(zph)\nplot(zph)')
add_tip(guide,'PH 不满足怎么办','确认数据和 functional form；考虑 time-by-covariate interaction、stratified Cox 或分时段 HR。报告“evidence of non-proportionality”，不要只删除变量。','FCE4D6')

guide.add_heading('9. Interaction / effect modification',level=1)
guide.add_paragraph('“Does the effect differ by...” 对应 PFT × modifier。正确检验是 interaction term 或嵌套模型 LRT，不是“一个 subgroup 显著、另一个不显著”。')
add_code(guide,'Surv(time, event) ~ PFT * modifier + covariates')
guide.add_paragraph('Interaction model 中 PFT main effect 是 modifier reference level 下的 PFT effect，不是全体平均 effect。若 interaction 有证据，应给 subgroup-specific estimates 或 adjusted predictions。')

guide.add_heading('10. Postoperative complications：PFT comparison + logistic regression',level=1)
guide.add_paragraph('第一步按 complication status 比较 PFT distribution。第二步对 any complication 建四个 separate adjusted logistic models。OR 必须解释为 odds，不应写成 risk，尤其当 complication 并不罕见时。')
add_code(guide,'glm(anycomp_bin ~ PFT + age + sex + smoking + type_of_surgery +\n      pathologic_stage + tumor_size_cm, family = binomial, data = model_dat)')
add_tip(guide,'Sparse outcome','death complication 等罕见结局可能出现 separation、极宽 CI 或不收敛。事件少于约 10 时，不要硬塞多变量模型；可保留描述性结果，并说明 precision limited。','FFF2CC')

guide.add_heading('11. Length of stay：不要轻易二分类',level=1)
guide.add_paragraph('LOS 通常右偏。先用 histogram、median/IQR 描述，再用 log(LOS) linear model；exp(beta) 解释为 geometric mean ratio。Gamma log-link 是 sensitivity analysis。除非有明确临床阈值，不建议仅为方便而分成 prolonged vs not prolonged。')
add_code(guide,'lm(log(length_of_hospital_stay) ~ PFT + covariates, data = los_dat)\n# exp(beta) = ratio of geometric mean LOS')
add_code(guide,'glm(length_of_hospital_stay ~ PFT + covariates,\n    family = Gamma(link = "log"), data = los_dat)')

guide.add_heading('12. Model building 与 N=130 的过拟合风险',level=1)
guide.add_paragraph('N=130 并不意味着可以放 20 个 parameters。真正限制来自 event count 和 sparse factor levels。脚本采用较小的 clinically prespecified core set，并让四个 PFT 模型保持相同 adjustment structure。')
for s in ['不要只按 univariate P<0.05 选择 confounders。','不要把 mediator 或 collider 当普通 confounder。','不要机械 stepwise 后把 selected model 当唯一真模型。','报告 model-specific N、事件数、CI width。','必要时把 full clinical model 放 sensitivity appendix。']:
    guide.add_paragraph(s,style='List Bullet')

guide.add_heading('13. 三线表 Word 输出',level=1)
guide.add_paragraph('脚本不用 gtsummary 的复杂 label pipeline，而是先整理普通 data frame，再用 flextable::theme_booktabs() 和 flextable::body_add_flextable()。这避免了你之前遇到的 modify_caption、Fisher workspace 和 label 类型错误。')
add_code(guide,'ft <- flextable::flextable(result_df) %>%\n  flextable::theme_booktabs() %>%\n  flextable::font(fontname = "Arial", part = "all")\n\ndoc <- officer::read_docx()\ndoc <- flextable::body_add_flextable(doc, value = ft)\nprint(doc, target = "output/word/results.docx")')

guide.add_heading('14. 每个 block 失败时怎么定位',level=1)
troubles=[('object not found','通常是前一块失败；从第一个 error 开始，不要继续往下跑。'),('unused argument in select','函数遮蔽；写 dplyr::select()。'),('number of rows has changed','不同模型 missingness 不一致；先创建共同 complete-case data。'),('model did not converge','检查 sparse cells、separation、空 factor level、过多 parameters。'),('cox.zph significant','检查 PH plot；考虑 time-varying effect 或 stratification。'),('very wide CI','事件少、collinearity 或 separation；减少复杂度并诚实报告 precision。'),('body_add_flextable not exported','使用 flextable::body_add_flextable()，不是 officer::body_add_flextable()。')]
t=guide.add_table(rows=1,cols=2); t.rows[0].cells[0].text='Error / signal'; t.rows[0].cells[1].text='Action'
for a,b in troubles: r=t.add_row().cells; r[0].text=a; r[1].text=b
style_three_line_table(t)

guide.add_heading('15. Report 写作顺序',level=1)
for i,s in enumerate(['Introduction：研究背景 + 三个 objectives。','Methods：population/variables → descriptive → survival → complications → LOS → interactions → diagnostics → missing data/software。','Results：先 cohort/missingness，再 descriptive，再 primary survival，再 secondary outcomes，再 diagnostics/sensitivity。','Discussion：总结 magnitude 与 uncertainty；不因 retrospective association 写因果；说明 N、小事件数、missing、residual confounding。','Tables/Figures：正文只保留回答 objectives 的 outputs；diagnostics 放 appendix。'],1):
    guide.add_paragraph(f'{i}. {s}')

guide.add_heading('16. Oral defense 高频问题',level=1)
qs=[('为什么四个 PFT 不放一个模型？','高度相关、不同尺度且题目明确要求 separate models；同时放入会造成 multicollinearity 并改变 estimand。'),('为什么 LOS 不直接 logistic？','二分类会损失信息且阈值可能任意；保持连续并处理 skewness 更符合 estimand。'),('为什么 HR 不是 RR？','HR 比较瞬时 hazard，依赖 survival process 和 PH interpretation。'),('为什么不按 univariate P 选 covariate？','confounding 是因果/设计概念，不等同于单变量显著性。'),('如果 PH 不满足？','time-varying coefficient、stratified Cox、分时段 effect 或替代 survival model。'),('为什么每个模型 N 不一样？','不同 PFT 与 covariate missingness；complete-case set 是 model-specific。')]
t=guide.add_table(rows=1,cols=2); t.rows[0].cells[0].text='Question'; t.rows[0].cells[1].text='Strong answer'
for a,b in qs: r=t.add_row().cells; r[0].text=a; r[1].text=b
style_three_line_table(t)

guide.add_heading('17. 考试当天 15 分钟导航',level=1)
steps=['抄写 3 个 aims；圈出 survival、binary complications、LOS。','建立 variable map；确认 event=1 的含义。','清理 factor/reference/missing；输出 cohort flow。','先做 Table 1/2/3 与 distribution plots。','跑 overall KM + univariate Cox。','跑四个 separate adjusted Cox；保存 HR/CI/P/N。','检查 PH；再做 interaction screening。','做 complication screening + anycomp logistic。','检查 LOS distribution；跑 log-linear + sensitivity。','停止探索，整理 Word outputs 并写 report。']
for i,s in enumerate(steps,1): guide.add_paragraph(f'{i}. {s}')
add_tip(guide,'最终原则','代码能跑只是最低标准。真正的合格答案必须让研究问题、estimand、模型、effect measure、diagnostics、table 和文字解释完全一致。')

guide_path=ROOT/'docs/QE2022_Applied_R_Code_and_Report_Guide_CN.docx'
guide.save(guide_path)

# Blank report template
report=Document(); setup_doc(report,'QE2022 Applied Practice Report Template','Lung cancer resection, pulmonary function, survival, complications, and length of stay')
for sec in report.sections: add_page_number(sec)
for heading,prompt in [
('Introduction','One concise paragraph: retrospective cohort, lung cancer resection, PFT exposures, three study objectives. Do not report results here.'),
('Statistical Methods','Describe cohort definition; variable coding; separate PFT models; descriptive statistics; KM/Cox; complication analyses; LOS strategy; interactions; diagnostics; missing-data handling; software/version.'),
('Results','Start with cohort and missingness. Then descriptive findings, survival, complications, LOS, interactions, and diagnostics. Report estimate + 95% CI + P value.'),
('Discussion','Summarize principal findings, clinical/epidemiologic interpretation, limitations, and cautious conclusion. Avoid causal language.'),
('References','List R package documentation, textbooks, and any external sources actually used.'),
('Tables and Figures','Insert publication-quality main tables/figures only.'),
('Appendix','Insert diagnostics, sensitivity analyses, and additional tables.')]:
    report.add_heading(heading,level=1); p=report.add_paragraph(prompt); p.runs[0].italic=True; p.runs[0].font.color.rgb=RGBColor(100,100,100)
report.save(ROOT/'docs/QE2022_Applied_Report_Template.docx')

# Summary file
summary = df.describe(include='all').transpose()
summary.to_csv(ROOT/'data/dictionary/synthetic_data_summary.csv')

# zip
zip_path=Path('/mnt/data/QE2022_Applied_Lung_Cancer_PFT_R_Project_with_Synthetic_Data.zip')
if zip_path.exists(): zip_path.unlink()
with zipfile.ZipFile(zip_path,'w',zipfile.ZIP_DEFLATED) as z:
    for fp in ROOT.rglob('*'):
        if fp.is_file(): z.write(fp, arcname=str(Path(ROOT.name)/fp.relative_to(ROOT)))

print(ROOT)
print(zip_path)
print(df.shape, int(df.died.sum()), int((df.anycomp=='Yes').sum()), df.Length_of_Hospital_stay.median())
